"""
Contract Template Service
Loads and processes contract templates from the contracts folder.
Migrated to work with FastAPI.
"""

from pathlib import Path
from typing import List, Dict, Any, Optional
from docx import Document

from app.config import get_settings


# Category mapping: display name -> folder name
CATEGORY_MAPPING = {
    "Аренда": "аренда",
    "Безвозмездное пользование": "безвозмедное пользование",
    "Дарение": "дарение",
    "Займ (кредит)": "Займ (кредит)",
    "Залог": "залог",
    "Купля-продажа, поставка, контрактация": "Купля-продажа, поставка, контрактация",
    "Подряд": "подряд",
    "Страхование": "страхование",
    "Услуги": "услуги",
    "Гражданско-правовой (ГПХ) договор": "gph_contract",
}

# Category icons
CATEGORY_ICONS = {
    "Аренда": "🏠",
    "Безвозмездное пользование": "🎁",
    "Дарение": "💝",
    "Займ (кредит)": "💰",
    "Залог": "🔒",
    "Купля-продажа, поставка, контрактация": "🛒",
    "Подряд": "🔨",
    "Страхование": "🛡️",
    "Услуги": "⚙️",
    "Гражданско-правовой (ГПХ) договор": "📄",
}


class ContractService:
    """Service for loading and processing contract templates."""
    
    def __init__(self, contracts_dir: Optional[str] = None):
        """Initialize with contracts directory path."""
        settings = get_settings()
        if contracts_dir:
            self.contracts_dir = Path(contracts_dir)
        else:
            self.contracts_dir = Path(settings.contracts_path)
    
    def get_categories(self) -> List[Dict[str, Any]]:
        """Get all contract categories with template counts."""
        categories = []
        
        for display_name, folder_name in CATEGORY_MAPPING.items():
            folder_path = self.contracts_dir / folder_name
            
            if folder_path.exists() and folder_path.is_dir():
                templates = list(folder_path.glob("*.docx"))
                template_count = len(templates)
                
                categories.append({
                    "name": display_name,
                    "count": template_count,
                    "description": CATEGORY_ICONS.get(display_name, "📄"),
                })
        
        return categories
    
    def get_templates_in_category(self, category: str) -> List[Dict[str, str]]:
        """Get list of templates in a category."""
        folder_name = CATEGORY_MAPPING.get(category)
        if not folder_name:
            return []
        
        folder_path = self.contracts_dir / folder_name
        if not folder_path.exists():
            return []
        
        templates = []
        for docx_file in folder_path.glob("*.docx"):
            templates.append({
                "name": docx_file.stem,
                "path": str(docx_file),
            })
        
        return templates
    
    def load_all_templates_for_category(self, category: str) -> str:
        """Load all templates for a category as combined context."""
        folder_name = CATEGORY_MAPPING.get(category)
        if not folder_name:
            return ""
        
        folder_path = self.contracts_dir / folder_name
        if not folder_path.exists():
            return ""
        
        all_content = []
        
        for docx_file in folder_path.glob("*.docx"):
            content = self._extract_docx_text(docx_file)
            if content:
                template_name = docx_file.stem
                all_content.append(f"=== ШАБЛОН: {template_name} ===\n\n{content}")
        
        return "\n\n" + "=" * 50 + "\n\n".join(all_content)
    
    def _extract_docx_text(self, docx_path: Path) -> Optional[str]:
        """Extract text content from a .docx file."""
        try:
            doc = Document(str(docx_path))
            paragraphs = []
            
            for para in doc.paragraphs:
                text = para.text.strip()
                if text:
                    paragraphs.append(text)
            
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        if cell_text:
                            row_text.append(cell_text)
                    if row_text:
                        paragraphs.append(" | ".join(row_text))
            
            return "\n\n".join(paragraphs)
            
        except Exception as e:
            print(f"Error extracting text from {docx_path}: {e}")
            return None
