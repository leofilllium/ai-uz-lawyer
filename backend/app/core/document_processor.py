"""
Document Processor Module
Handles extraction, parsing, and chunking of legal documents from DOCX files.
Migrated to work with FastAPI.
"""

import re
import uuid
from pathlib import Path
from typing import List, Dict, Any, Optional, Tuple
from docx import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter

from app.config import get_settings


class DocumentProcessor:
    """Process legal documents from DOCX files into searchable chunks."""
    
    def __init__(self, chunk_size: int = 1500, chunk_overlap: int = 200):
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
            separators=["\n\n\n", "\n\n", "\n", ". ", ", ", " "],
            length_function=len,
        )
        self._article_display_map: Dict[int, str] = {}
    
    def extract_text_from_docx(self, file_path: Path) -> str:
        """Extract all text content from a DOCX file."""
        doc = Document(file_path)
        full_text = []
        
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if text:
                full_text.append(text)
        
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text = cell.text.strip()
                    if text:
                        full_text.append(text)
        
        return "\n\n".join(full_text)
    
    def _build_article_display_map(self, text: str) -> Dict[int, str]:
        """Build a map from article positions to their display names."""
        article_pattern = re.compile(r'Статья\s+(\d+)', re.IGNORECASE)
        
        article_list: List[Tuple[int, int]] = []
        for match in article_pattern.finditer(text):
            try:
                num = int(match.group(1))
                article_list.append((match.start(), num))
            except ValueError:
                pass
        
        display_map: Dict[int, str] = {}
        
        for i, (pos, num) in enumerate(article_list):
            prev_num = article_list[i-1][1] if i > 0 else None
            next_num = article_list[i+1][1] if i < len(article_list)-1 else None
            
            display = str(num)
            
            if num < 100:
                display_map[pos] = display
                continue
            
            if num % 10 == 0:
                display_map[pos] = display
                continue
            
            base = num // 10
            sub = num % 10
            
            if num >= 1000:
                display = f"{base}.{sub}"
                display_map[pos] = display
                continue
            
            if prev_num == num - 1 or next_num == num + 1:
                display_map[pos] = display
                continue
            
            if prev_num == base and next_num == base + 1:
                display = f"{base}.{sub}"
                display_map[pos] = display
                continue
            
            if prev_num is not None:
                prev_base = prev_num // 10 if prev_num >= 100 else prev_num
                if prev_num == base or (prev_base == base and prev_num >= 100 and prev_num < 1000):
                    for j in range(i+1, min(i+10, len(article_list))):
                        future_num = article_list[j][1]
                        if future_num == base + 1:
                            display = f"{base}.{sub}"
                            break
                        if future_num == num + 1 or future_num > base + 1:
                            break
            
            display_map[pos] = display
        
        return display_map
    
    def parse_article_structure(self, text: str) -> List[Dict[str, Any]]:
        """Parse Russian/Uzbek legal document."""
        self._article_display_map = self._build_article_display_map(text)
        
        articles = []
        current_section = "ОБЩАЯ ЧАСТЬ"
        current_chapter = "General"
        current_chapter_num = ""
        current_article = None
        current_article_display = None
        current_article_title = ""
        current_text = []
        
        section_pattern = re.compile(r'^(РАЗДЕЛ\s+[А-ЯA-Z]+|ОБЩАЯ ЧАСТЬ|ОСОБЕННАЯ ЧАСТЬ)', re.IGNORECASE)
        chapter_pattern = re.compile(r'^Глава\s+([IVXLC]+|\d+)[.\s]*(.*)$', re.IGNORECASE)
        article_pattern = re.compile(r'^Статья\s+(\d+)[.\s]*(.*)$', re.IGNORECASE)
        
        current_pos = 0
        lines = text.split('\n')
        
        for line in lines:
            line_stripped = line.strip()
            if not line_stripped:
                current_pos += len(line) + 1
                continue
            
            section_match = section_pattern.match(line_stripped)
            if section_match:
                current_section = line_stripped
                current_pos += len(line) + 1
                continue
            
            chapter_match = chapter_pattern.match(line_stripped)
            if chapter_match:
                current_chapter_num = chapter_match.group(1)
                chapter_title = chapter_match.group(2).strip() if chapter_match.group(2) else ""
                current_chapter = f"Глава {current_chapter_num}. {chapter_title}".strip()
                current_pos += len(line) + 1
                continue
            
            article_match = article_pattern.match(line_stripped)
            if article_match:
                if current_article and current_text:
                    articles.append({
                        "article_number": current_article,
                        "article_display": current_article_display,
                        "chapter": current_chapter,
                        "chapter_num": current_chapter_num,
                        "section": current_section,
                        "title": current_article_title,
                        "content": "\n".join(current_text),
                    })
                
                raw_article_num = article_match.group(1)
                current_article = raw_article_num
                
                article_pos = text.find(f"Статья {raw_article_num}", current_pos - 100)
                if article_pos == -1:
                    article_pos = text.find(f"Статья {raw_article_num}")
                current_article_display = self._article_display_map.get(article_pos, raw_article_num)
                
                current_article_title = article_match.group(2).strip() if article_match.group(2) else ""
                current_text = [line_stripped]
                current_pos += len(line) + 1
                continue
            
            if current_article:
                current_text.append(line_stripped)
            current_pos += len(line) + 1
        
        if current_article and current_text:
            articles.append({
                "article_number": current_article,
                "article_display": current_article_display,
                "chapter": current_chapter,
                "chapter_num": current_chapter_num,
                "section": current_section,
                "title": current_article_title,
                "content": "\n".join(current_text),
            })
        
        return articles
    
    def create_chunks_with_metadata(self, text: str, source_file: str) -> List[Dict[str, Any]]:
        """Split text into chunks while preserving metadata."""
        chunks = []
        articles = self.parse_article_structure(text)
        
        if articles:
            for article in articles:
                article_chunks = self.text_splitter.split_text(article["content"])
                for i, chunk in enumerate(article_chunks):
                    chunks.append({
                        "id": str(uuid.uuid4()),
                        "content": chunk,
                        "metadata": {
                            "source": source_file,
                            "article_number": article["article_number"],
                            "article_display": article["article_display"],
                            "chapter": article["chapter"][:200] if article["chapter"] else "",
                            "chapter_num": article.get("chapter_num", ""),
                            "section": article.get("section", "")[:100],
                            "title": article["title"][:150] if article["title"] else "",
                            "chunk_index": i,
                            "total_chunks": len(article_chunks),
                        }
                    })
        else:
            text_chunks = self.text_splitter.split_text(text)
            for i, chunk in enumerate(text_chunks):
                chunks.append({
                    "id": str(uuid.uuid4()),
                    "content": chunk,
                    "metadata": {
                        "source": source_file,
                        "article_number": "Unknown",
                        "article_display": "Unknown",
                        "chapter": "General",
                        "chunk_index": i,
                        "total_chunks": len(text_chunks),
                    }
                })
        
        return chunks
    
    def process_documents(self, codes_directory: Optional[Path] = None) -> List[Dict[str, Any]]:
        """Process all DOCX files in the codes directory."""
        settings = get_settings()
        
        if codes_directory is None:
            codes_directory = Path(settings.codes_path)
        else:
            codes_directory = Path(codes_directory)
        
        all_chunks = []
        docx_files = list(codes_directory.glob("*.docx"))
        
        if not docx_files:
            print(f"No DOCX files found in {codes_directory}")
            return []
        
        for docx_file in docx_files:
            print(f"Processing: {docx_file.name}")
            text = self.extract_text_from_docx(docx_file)
            chunks = self.create_chunks_with_metadata(text, docx_file.name)
            all_chunks.extend(chunks)
            print(f"  {len(chunks)} chunks created")
        
        print(f"\nTotal: {len(all_chunks)} chunks")
        return all_chunks
