"""
Flexible Document Processor
Handles extraction, parsing, and chunking of legal documents from DOCX files.
Supports both Russian and Uzbek legal document formats.
"""

import re
import uuid
import hashlib
from pathlib import Path
from typing import List, Dict, Any, Optional, Tuple
from docx import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter


class FlexibleDocumentProcessor:
    """Process legal documents from both Russian and Uzbek formats."""
    
    # Document type constants
    TYPE_RUSSIAN_CODE = "russian_code"
    TYPE_UZBEK_CODE = "uzbek_code"
    TYPE_DECREE = "decree"
    TYPE_GENERIC = "generic"
    
    def __init__(self, chunk_size: int = 1500, chunk_overlap: int = 200):
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
            separators=["\n\n\n", "\n\n", "\n", ". ", ", ", " "],
            length_function=len,
        )
    
    def extract_text_from_docx(self, file_path: Path) -> str:
        """Extract all text content from a DOCX file."""
        doc = Document(file_path)
        full_text = []
        
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if text:
                full_text.append(text)
        
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text = cell.text.strip()
                    if text:
                        full_text.append(text)
        
        return "\n\n".join(full_text)
    
    def detect_document_type(self, text: str) -> str:
        """Detect the type/format of the legal document."""
        # Count pattern occurrences
        russian_article_count = len(re.findall(r'Статья\s+\d+', text, re.IGNORECASE))
        uzbek_article_count = len(re.findall(r'\d+-modda', text, re.IGNORECASE))
        
        # Check for decree patterns
        decree_patterns = [
            r'QAROR',
            r'NIZOM',
            r'QONUN',
            r'FARMON',
            r'QAROR\s+QILADI',
        ]
        decree_count = sum(1 for p in decree_patterns if re.search(p, text, re.IGNORECASE))
        
        # Determine type based on counts
        if uzbek_article_count >= 3:
            return self.TYPE_UZBEK_CODE
        elif russian_article_count >= 3:
            return self.TYPE_RUSSIAN_CODE
        elif decree_count >= 1:
            return self.TYPE_DECREE
        else:
            return self.TYPE_GENERIC
    
    def parse_russian_structure(self, text: str) -> List[Dict[str, Any]]:
        """Parse Russian legal document structure (Статья, Глава, Раздел)."""
        articles = []
        current_section = "ОБЩАЯ ЧАСТЬ"
        current_chapter = "General"
        current_chapter_num = ""
        current_article = None
        current_article_title = ""
        current_text = []
        
        section_pattern = re.compile(r'^(РАЗДЕЛ\s+[А-ЯA-Z]+|ОБЩАЯ ЧАСТЬ|ОСОБЕННАЯ ЧАСТЬ)', re.IGNORECASE)
        chapter_pattern = re.compile(r'^Глава\s+([IVXLC]+|\d+)[.\s]*(.*)$', re.IGNORECASE)
        article_pattern = re.compile(r'^Статья\s+(\d+)[.\s]*(.*)$', re.IGNORECASE)
        
        lines = text.split('\n')
        
        for line in lines:
            line_stripped = line.strip()
            if not line_stripped:
                continue
            
            section_match = section_pattern.match(line_stripped)
            if section_match:
                current_section = line_stripped
                continue
            
            chapter_match = chapter_pattern.match(line_stripped)
            if chapter_match:
                current_chapter_num = chapter_match.group(1)
                chapter_title = chapter_match.group(2).strip() if chapter_match.group(2) else ""
                current_chapter = f"Глава {current_chapter_num}. {chapter_title}".strip()
                continue
            
            article_match = article_pattern.match(line_stripped)
            if article_match:
                # Save previous article
                if current_article and current_text:
                    articles.append({
                        "article_number": current_article,
                        "article_display": current_article,
                        "chapter": current_chapter,
                        "chapter_num": current_chapter_num,
                        "section": current_section,
                        "title": current_article_title,
                        "content": "\n".join(current_text),
                    })
                
                current_article = article_match.group(1)
                current_article_title = article_match.group(2).strip() if article_match.group(2) else ""
                current_text = [line_stripped]
                continue
            
            if current_article:
                current_text.append(line_stripped)
        
        # Save last article
        if current_article and current_text:
            articles.append({
                "article_number": current_article,
                "article_display": current_article,
                "chapter": current_chapter,
                "chapter_num": current_chapter_num,
                "section": current_section,
                "title": current_article_title,
                "content": "\n".join(current_text),
            })
        
        return articles
    
    def parse_uzbek_structure(self, text: str) -> List[Dict[str, Any]]:
        """Parse Uzbek legal document structure (X-modda, X bob, BOʻLIM)."""
        articles = []
        current_section = "UMUMIY QOIDALAR"
        current_chapter = "General"
        current_chapter_num = ""
        current_article = None
        current_article_title = ""
        current_text = []
        
        # Uzbek patterns
        # BOʻLIM (section), bob (chapter), modda (article)
        section_pattern = re.compile(r'^(BIRINCHI|IKKINCHI|UCHINCHI|TOʻRTINCHI|BESHINCHI)?\s*BOʻLIM', re.IGNORECASE)
        chapter_pattern = re.compile(r'^([IVXLC]+|\d+)\s*bob[.\s]*(.*)$', re.IGNORECASE)
        article_pattern = re.compile(r'^(\d+)-modda[.\s]*(.*)$', re.IGNORECASE)
        
        lines = text.split('\n')
        
        for line in lines:
            line_stripped = line.strip()
            if not line_stripped:
                continue
            
            section_match = section_pattern.match(line_stripped)
            if section_match:
                current_section = line_stripped
                continue
            
            chapter_match = chapter_pattern.match(line_stripped)
            if chapter_match:
                current_chapter_num = chapter_match.group(1)
                chapter_title = chapter_match.group(2).strip() if chapter_match.group(2) else ""
                current_chapter = f"{current_chapter_num} bob. {chapter_title}".strip()
                continue
            
            article_match = article_pattern.match(line_stripped)
            if article_match:
                # Save previous article
                if current_article and current_text:
                    articles.append({
                        "article_number": current_article,
                        "article_display": f"{current_article}-modda",
                        "chapter": current_chapter,
                        "chapter_num": current_chapter_num,
                        "section": current_section,
                        "title": current_article_title,
                        "content": "\n".join(current_text),
                    })
                
                current_article = article_match.group(1)
                current_article_title = article_match.group(2).strip() if article_match.group(2) else ""
                current_text = [line_stripped]
                continue
            
            if current_article:
                current_text.append(line_stripped)
        
        # Save last article
        if current_article and current_text:
            articles.append({
                "article_number": current_article,
                "article_display": f"{current_article}-modda",
                "chapter": current_chapter,
                "chapter_num": current_chapter_num,
                "section": current_section,
                "title": current_article_title,
                "content": "\n".join(current_text),
            })
        
        return articles
    
    def parse_decree_structure(self, text: str) -> List[Dict[str, Any]]:
        """Parse government decrees, laws, and regulations (QAROR, NIZOM, QONUN)."""
        articles = []
        
        # Try to extract title
        title_match = re.search(r'(QONUNI?|QARORI?|FARMONI?|NIZOMI?)\s*\n(.+?)(?=\n\n|\n\d+\.)', text, re.IGNORECASE | re.DOTALL)
        document_title = title_match.group(2).strip()[:200] if title_match else "Unknown Document"
        
        # Split by numbered points (1., 2., 3., etc.)
        point_pattern = re.compile(r'^(\d+)\.\s+(.+?)(?=^\d+\.\s|\Z)', re.MULTILINE | re.DOTALL)
        
        matches = list(point_pattern.finditer(text))
        
        if matches:
            for match in matches:
                point_num = match.group(1)
                content = match.group(2).strip()
                
                articles.append({
                    "article_number": point_num,
                    "article_display": f"пункт {point_num}",
                    "chapter": "General",
                    "chapter_num": "",
                    "section": document_title[:100],
                    "title": content[:100] + "..." if len(content) > 100 else content,
                    "content": content,
                })
        
        # Also try to parse NIZOM sections if present
        nizom_pattern = re.compile(r'^(\d+)\.\s+"?([^"]+)"?\s+(ordeni|toʻgʻrisidagi|haqida)', re.IGNORECASE | re.MULTILINE)
        
        # If we found very few points, also try splitting by paragraphs
        if len(articles) < 3:
            # Split into meaningful paragraphs
            paragraphs = [p.strip() for p in text.split('\n\n') if len(p.strip()) > 100]
            
            for i, para in enumerate(paragraphs):
                # Skip if already covered
                if any(para[:50] in a.get("content", "") for a in articles):
                    continue
                
                articles.append({
                    "article_number": str(i + 1),
                    "article_display": f"раздел {i + 1}",
                    "chapter": "General",
                    "chapter_num": "",
                    "section": document_title[:100],
                    "title": para[:80] + "..." if len(para) > 80 else para,
                    "content": para,
                })
        
        return articles
    
    def process_single_document(
        self, 
        file_path: Path, 
        source_name: Optional[str] = None
    ) -> Tuple[List[Dict[str, Any]], Dict[str, Any]]:
        """
        Process a single document and return chunks with metadata.
        
        Returns:
            Tuple of (chunks, document_info)
        """
        if source_name is None:
            source_name = file_path.name
        
        # Extract text
        text = self.extract_text_from_docx(file_path)
        
        # Generate document hash for tracking
        doc_hash = hashlib.md5(text.encode()).hexdigest()[:12]
        
        # Detect document type
        doc_type = self.detect_document_type(text)
        
        # Parse based on type
        if doc_type == self.TYPE_RUSSIAN_CODE:
            articles = self.parse_russian_structure(text)
        elif doc_type == self.TYPE_UZBEK_CODE:
            articles = self.parse_uzbek_structure(text)
        elif doc_type == self.TYPE_DECREE:
            articles = self.parse_decree_structure(text)
        else:
            articles = []
        
        # Create chunks
        chunks = []
        
        if articles:
            for article in articles:
                article_chunks = self.text_splitter.split_text(article["content"])
                for i, chunk in enumerate(article_chunks):
                    chunks.append({
                        "id": str(uuid.uuid4()),
                        "content": chunk,
                        "metadata": {
                            "source": source_name,
                            "article_number": str(article["article_number"]),
                            "article_display": str(article["article_display"]),
                            "chapter": article["chapter"][:200] if article["chapter"] else "",
                            "chapter_num": article.get("chapter_num", ""),
                            "section": article.get("section", "")[:100],
                            "title": article["title"][:150] if article["title"] else "",
                            "chunk_index": i,
                            "total_chunks": len(article_chunks),
                            "doc_type": doc_type,
                        }
                    })
        else:
            # Fallback: split entire text
            text_chunks = self.text_splitter.split_text(text)
            for i, chunk in enumerate(text_chunks):
                chunks.append({
                    "id": str(uuid.uuid4()),
                    "content": chunk,
                    "metadata": {
                        "source": source_name,
                        "article_number": "Unknown",
                        "article_display": "Unknown",
                        "chapter": "General",
                        "chapter_num": "",
                        "section": "",
                        "title": "",
                        "chunk_index": i,
                        "total_chunks": len(text_chunks),
                        "doc_type": doc_type,
                    }
                })
        
        # Document info for tracking
        doc_info = {
            "source_name": source_name,
            "doc_hash": doc_hash,
            "doc_type": doc_type,
            "article_count": len(articles),
            "chunk_count": len(chunks),
            "text_length": len(text),
        }
        
        return chunks, doc_info
    
    def process_documents(self, codes_directory: Optional[Path] = None) -> List[Dict[str, Any]]:
        """Process all DOCX files in a directory (for bulk indexing)."""
        from app.config import get_settings
        settings = get_settings()
        
        if codes_directory is None:
            codes_directory = Path(settings.codes_path)
        else:
            codes_directory = Path(codes_directory)
        
        all_chunks = []
        docx_files = list(codes_directory.glob("*.docx"))
        
        if not docx_files:
            print(f"No DOCX files found in {codes_directory}")
            return []
        
        for docx_file in docx_files:
            print(f"Processing: {docx_file.name}")
            chunks, doc_info = self.process_single_document(docx_file)
            all_chunks.extend(chunks)
            print(f"  Type: {doc_info['doc_type']}, Articles: {doc_info['article_count']}, Chunks: {doc_info['chunk_count']}")
        
        print(f"\nTotal: {len(all_chunks)} chunks")
        return all_chunks
